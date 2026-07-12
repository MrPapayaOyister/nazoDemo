# -*- coding: utf-8 -*-
"""Real document generation for a correspondence — signed PDF + best-effort DOCX.

Pipeline (STEP 7):
  render_letter_html  — substitute EVERY {{TOKEN}} in the template's doc_html into
                        a self-contained, embeddable HTML letter (pure HTML/CSS
                        letterhead, inline signature <img> from the signature
                        data-URI, deterministic amount/date formatting) and wrap
                        it in a full A4 print document.
  render_pdf          — POST that HTML to Gotenberg's Chromium module -> PDF bytes.
  render_docx         — best-effort HTML -> DOCX (htmldocx / python-docx).
  snapshot_version    — persist a CorrespondenceVersion row (rendered html + values
                        + pdf/docx bytes) for audit on approval.

Design notes:
  * We render the EXISTING HTML (template.doc_html carries the letter) — we do NOT
    use docxtpl (there is no .docx letterhead; the letterhead is pure HTML/CSS).
  * The on-screen document (nazo-ai DocumentRenderer + Letterhead) is mirrored so
    the printed letter looks like the app's document (same crest, rule, sign-block).
  * Arabic renders RTL with the Amiri / Noto Naskh Arabic fonts installed in the
    Gotenberg image; Western digits + Gregorian dates come from app.services.amounts.
  * User-supplied values are HTML-escaped (no injection). Only the intended
    signature <img> and the letterhead block are emitted as raw markup.
"""

from __future__ import annotations

import logging
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional

import httpx
from sqlalchemy.exc import IntegrityError
from sqlmodel import Session, select

from app.config import settings
from app.db import engine
from app.models import (
    AppUser,
    Correspondence,
    CorrespondenceVersion,
    Signature,
    Template,
)
from babel.numbers import format_decimal

from app.services.amounts import format_date, group_number

logger = logging.getLogger("nazo.documents")

# Mirror of the frontend {{TOKEN}} matcher (src/components/common/DocumentRenderer).
_TOKEN_RE = re.compile(r"\{\{\s*([A-Z0-9_]+)\s*\}\}")

# Brand colours (mirror nazo-ai globals.css --navy / --gold).
_NAVY = "#12336b"
_GOLD = "#b0871c"

# ORG constants (mirror nazo-ai/src/lib/constants.ts ORG).
_ORG = {
    "nameEn": "Education, Human Development & Community Development Council",
    "nameAr": "مجلس التعليم والتنمية البشرية والتنمية المجتمعية",
    "subEn": "Federal Authority for Government Human Resources",
    "subAr": "الهيئة الاتحادية للموارد البشرية الحكومية",
    "cityEn": "Abu Dhabi, United Arab Emirates",
    "cityAr": "أبوظبي، الإمارات العربية المتحدة",
    "poBox": "P.O. Box 33845",
    "web": "www.ehcd.gov.ae",
}


# ---------------------------------------------------------------------------
# Small local helpers (kept independent of workflow.py to avoid coupling).
# ---------------------------------------------------------------------------
def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _gen_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


def escape_html(s: str) -> str:
    """Escape the five HTML-significant characters (never trust user values)."""
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def _resolve_lang(template: Optional[Template], lang: Optional[str]) -> str:
    if lang in ("en", "ar"):
        return lang  # explicit override wins
    if template is not None and template.lang in ("en", "ar"):
        return template.lang
    return "en"


def _format_amount(value: str) -> Optional[str]:
    """Grouped Western number for an amount-like string, or None if not numeric.

    Strips grouping commas/spaces. Integers -> '75,000'; decimals are PRESERVED
    (no cent-dropping): '75000.50' -> '75,000.5', '1234.75' -> '1,234.75'. Returns
    None when the value isn't purely numeric (caller falls back to the raw text).
    """
    cleaned = re.sub(r"[,\s]", "", value or "")
    if not re.fullmatch(r"\d+(\.\d+)?", cleaned):
        return None
    if "." in cleaned:
        # Non-integer: format via babel float path so decimals survive.
        return format_decimal(float(cleaned), locale="en_US")
    return group_number(int(cleaned))


# ---------------------------------------------------------------------------
# Letterhead (pure HTML/CSS — no external image assets; fully embeddable).
# ---------------------------------------------------------------------------
def _letterhead_html(lang: str) -> str:
    is_ar = lang == "ar"
    dir_attr = "rtl" if is_ar else "ltr"
    org = _ORG["nameAr"] if is_ar else _ORG["nameEn"]
    sub = _ORG["subAr"] if is_ar else _ORG["subEn"]
    city = _ORG["cityAr"] if is_ar else _ORG["cityEn"]
    crest = (
        '<svg viewBox="0 0 48 48" width="46" height="46" '
        'xmlns="http://www.w3.org/2000/svg">'
        f'<rect x="1" y="1" width="46" height="46" rx="9" fill="{_NAVY}"/>'
        f'<path d="M24 9 L37 15 V25 C37 33 31 38 24 40 C17 38 11 33 11 25 V15 Z" '
        f'fill="none" stroke="{_GOLD}" stroke-width="2"/>'
        '<text x="24" y="29" text-anchor="middle" font-size="15" '
        'font-weight="700" fill="#fff" font-family="serif">E</text></svg>'
    )
    return (
        f'<div class="doc-letterhead" dir="{dir_attr}">'
        '<div class="doc-lh-row">'
        f'<div class="doc-crest">{crest}</div>'
        '<div class="doc-lh-titles">'
        f'<div class="doc-lh-org">{escape_html(org)}</div>'
        f'<div class="doc-lh-sub">{escape_html(sub)}</div>'
        "</div>"
        '<div class="doc-lh-meta">'
        f'<div>{escape_html(_ORG["poBox"])}</div>'
        f'<div>{escape_html(city)}</div>'
        f'<div>{escape_html(_ORG["web"])}</div>'
        "</div>"
        "</div>"
        '<div class="doc-lh-rule"></div>'
        "</div>"
    )


# ---------------------------------------------------------------------------
# Signature block (inline <img> from the SVG data-URI).
# ---------------------------------------------------------------------------
def _signature_html(
    session: Session, signature_id: str, lang: str, *, for_docx: bool
) -> str:
    sig = session.get(Signature, signature_id)
    if sig is None:
        return _empty_sig_html(lang)
    owner = session.get(AppUser, sig.owner_id)
    is_ar = lang == "ar"
    name = ""
    title = ""
    if owner is not None:
        name = owner.name_ar if is_ar else owner.name_en
        title = owner.title_ar if is_ar else owner.title_en

    if for_docx:
        # SVG data-URIs cannot be embedded by python-docx; caption the signer.
        label = "موقّع" if is_ar else "Signed"
        return (
            f'<p><strong>{escape_html(name)}</strong><br/>'
            f'{escape_html(title)}<br/>'
            f'<em>({escape_html(label)})</em></p>'
        )

    # data_uri is app-generated (seed) SVG markup, but defensively neutralize any
    # literal double-quote so it can never break out of the src="" attribute
    # context (correctness no longer relies on the data_uri being pre-encoded).
    src = (sig.data_uri or "").replace('"', "%22")
    return (
        '<span class="doc-sig">'
        f'<img class="sig doc-sig-img" src="{src}" alt="signature"/>'
        '<span class="doc-sig-cap">'
        f"{escape_html(name)}<br/>"
        f'<span class="doc-sig-role">{escape_html(title)}</span>'
        "</span>"
        "</span>"
    )


def _empty_sig_html(lang: str) -> str:
    label = "التوقيع" if lang == "ar" else "Signature"
    return (
        '<span class="doc-sig doc-sig--empty">'
        f'<span class="doc-sig-slot">{escape_html(label)}</span>'
        "</span>"
    )


# ---------------------------------------------------------------------------
# Token substitution over the template body.
# ---------------------------------------------------------------------------
def _substitute_body(
    session: Session,
    template: Template,
    values: dict[str, str],
    lang: str,
    *,
    for_docx: bool,
) -> str:
    sig_tags = {
        v["tag"]
        for v in (template.variables or [])
        if isinstance(v, dict) and v.get("type") == "Signature"
    }

    def repl(match: re.Match) -> str:
        name = match.group(1)
        tag = f"{{{{{name}}}}}"

        if name == "LETTERHEAD":
            if for_docx:
                org = _ORG["nameAr"] if lang == "ar" else _ORG["nameEn"]
                sub = _ORG["subAr"] if lang == "ar" else _ORG["subEn"]
                return f"<h2>{escape_html(org)}</h2><p>{escape_html(sub)}</p><hr/>"
            return _letterhead_html(lang)

        is_sig = tag in sig_tags or name.startswith("SIG")
        raw = values.get(tag, "")

        if is_sig:
            if raw:
                return _signature_html(session, raw, lang, for_docx=for_docx)
            if for_docx:
                return ""
            return _empty_sig_html(lang)

        if raw == "" or raw is None:
            # Unfilled token -> subtle blank underline (nothing to leak).
            if for_docx:
                return "____________"
            return '<span class="blank"></span>'

        # Amount-like -> deterministic grouped number when numeric (decimals kept).
        if "AMOUNT" in name:
            grouped = _format_amount(raw)
            if grouped is not None:
                return escape_html(grouped)
            return escape_html(raw)

        # Date-like -> Gregorian, Western digits, localized month.
        if "DATE" in name:
            return escape_html(format_date(raw, lang))

        return escape_html(raw)

    return _TOKEN_RE.sub(repl, template.doc_html or "")


# ---------------------------------------------------------------------------
# Full HTML document (A4 print styling mirroring the app's letter).
# ---------------------------------------------------------------------------
def _document_css(lang: str) -> str:
    ar_font = (
        "html[dir='rtl'] body, html[dir='rtl'] .nazo-doc {"
        " font-family: 'Amiri','Noto Naskh Arabic','Scheherazade New',serif; }"
    )
    return f"""
@page {{ size: A4; margin: 18mm 16mm; }}
* {{ box-sizing: border-box; }}
html, body {{ margin: 0; padding: 0; background: #ffffff; }}
body {{
  color: #16233d;
  font-family: 'Inter','Segoe UI',Arial,sans-serif;
  font-size: 13px;
  line-height: 1.75;
  -webkit-print-color-adjust: exact;
  print-color-adjust: exact;
}}
{ar_font}
.nazo-doc {{ background: #ffffff; color: #16233d; }}
.doc-letterhead {{ margin-bottom: 22px; }}
.doc-lh-row {{ display: flex; align-items: center; gap: 14px; }}
.doc-crest {{ flex-shrink: 0; line-height: 0; }}
.doc-lh-titles {{ flex: 1; min-width: 0; }}
.doc-lh-org {{ font-weight: 700; font-size: 15px; color: {_NAVY}; letter-spacing: -0.01em; }}
.doc-lh-sub {{ font-size: 11px; color: #5a6b8c; }}
.doc-lh-meta {{ text-align: end; font-size: 10.5px; color: #6b7a97; line-height: 1.5; flex-shrink: 0; }}
.doc-lh-rule {{ margin-top: 12px; height: 3px; border-radius: 3px;
  background: linear-gradient(90deg, {_NAVY} 0%, {_GOLD} 60%, transparent 100%); }}
html[dir='rtl'] .doc-lh-rule {{
  background: linear-gradient(270deg, {_NAVY} 0%, {_GOLD} 60%, transparent 100%); }}
h1 {{ font-size: 17px; font-weight: 700; color: {_NAVY}; margin: 6px 0 12px; }}
h2 {{ font-size: 15px; font-weight: 700; color: #16233d; margin: 14px 0 8px; }}
p {{ margin: 9px 0; }}
.meta {{ color: #4a5b7d; font-size: 13px; }}
.blank {{ display: inline-block; min-width: 90px; border-bottom: 1px solid #9aa8c2;
  height: 1em; vertical-align: bottom; }}
.sign-block {{ display: flex; gap: 34px; flex-wrap: wrap; margin-top: 34px; }}
.doc-sig {{ display: inline-flex; flex-direction: column; align-items: center; min-width: 150px; }}
.sig, .doc-sig-img {{ width: 140px; height: 52px; object-fit: contain; }}
.doc-sig-cap {{ margin-top: 4px; padding-top: 4px; border-top: 1px solid #d7deea;
  width: 150px; text-align: center; font-size: 11px; font-weight: 600; color: #24365a; }}
.doc-sig-role {{ font-weight: 500; color: #7183a3; font-size: 10px; }}
.doc-sig--empty .doc-sig-slot {{ display: grid; place-items: center; width: 140px;
  height: 52px; border: 1px dashed #c2cee0; border-radius: 8px; color: #9aa8c2; font-size: 11px; }}
""".strip()


def render_letter_html(
    session: Session, corr: Correspondence, *, lang: Optional[str] = None
) -> str:
    """Fully substituted, self-contained HTML letter for `corr` (A4 print doc)."""
    template = session.get(Template, corr.template_id)
    resolved = _resolve_lang(template, lang)
    if template is None:
        # Degrade rather than 500: emit a minimal but valid document.
        body = _letterhead_html(resolved) + "<p>Template not found.</p>"
    else:
        body = _substitute_body(
            session, template, dict(corr.values or {}), resolved, for_docx=False
        )
    dir_attr = "rtl" if resolved == "ar" else "ltr"
    return (
        "<!doctype html>"
        f'<html lang="{resolved}" dir="{dir_attr}">'
        '<head><meta charset="utf-8"/>'
        f"<style>{_document_css(resolved)}</style></head>"
        f'<body><div class="nazo-doc"><div class="doc-body">{body}</div></div></body>'
        "</html>"
    )


# ---------------------------------------------------------------------------
# PDF via Gotenberg Chromium module.
# ---------------------------------------------------------------------------
def render_pdf(
    session: Session, corr: Correspondence, *, lang: Optional[str] = None
) -> bytes:
    """Render the letter HTML to PDF bytes through Gotenberg's Chromium module."""
    html = render_letter_html(session, corr, lang=lang)
    url = f"{settings.gotenberg_url.rstrip('/')}/forms/chromium/convert/html"
    # Gotenberg REQUIRES the main file to be named exactly 'index.html'.
    files = {"index.html": ("index.html", html.encode("utf-8"), "text/html")}
    # Single source of truth for page geometry: the CSS '@page { size:A4; margin }'
    # in _document_css. preferCssPageSize lets Chromium honour that size AND margin,
    # so we do NOT also send Gotenberg marginTop/Bottom/Left/Right (they would be
    # a second, drift-prone definition).
    data = {
        "printBackground": "true",
        "preferCssPageSize": "true",
    }
    resp = httpx.post(url, files=files, data=data, timeout=30.0)
    if resp.status_code != 200:
        raise RuntimeError(
            f"Gotenberg PDF render failed ({resp.status_code}): "
            f"{resp.text[:400]}"
        )
    return resp.content


# ---------------------------------------------------------------------------
# DOCX (best-effort HTML -> DOCX; must never break the PDF path).
# ---------------------------------------------------------------------------
def render_docx(
    session: Session, corr: Correspondence, *, lang: Optional[str] = None
) -> bytes:
    """Best-effort DOCX of the letter. Prefers htmldocx; falls back to plain text."""
    template = session.get(Template, corr.template_id)
    resolved = _resolve_lang(template, lang)
    if template is None:
        body_html = "<p>Template not found.</p>"
    else:
        body_html = _substitute_body(
            session, template, dict(corr.values or {}), resolved, for_docx=True
        )

    # Primary path: htmldocx renders the simplified body into a python-docx doc.
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from htmldocx import HtmlToDocx  # type: ignore

        document = Document()
        if resolved == "ar":
            # Right-align the default paragraph style for RTL letters.
            try:
                document.styles["Normal"].paragraph_format.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                )
            except Exception:  # noqa: BLE001
                pass
        parser = HtmlToDocx()
        parser.add_html_to_document(f"<div>{body_html}</div>", document)
        bio = BytesIO()
        document.save(bio)
        return bio.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.warning("htmldocx DOCX render failed, falling back to text: %s", exc)

    # Fallback: minimal python-docx from stripped text.
    try:
        from docx import Document  # type: ignore

        text = _strip_tags(body_html)
        document = Document()
        for line in [ln.strip() for ln in text.splitlines()]:
            if line:
                document.add_paragraph(line)
        if not document.paragraphs:
            document.add_paragraph(corr.ref or "")
        bio = BytesIO()
        document.save(bio)
        return bio.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.error("Fallback DOCX render failed: %s", exc)
        raise RuntimeError(f"DOCX render failed: {exc}") from exc


def _strip_tags(html: str) -> str:
    """Best-effort tag stripping (BeautifulSoup if present, else regex)."""
    try:
        from bs4 import BeautifulSoup  # type: ignore

        return BeautifulSoup(html, "html.parser").get_text("\n")
    except Exception:  # noqa: BLE001
        text = re.sub(r"<br\s*/?>", "\n", html)
        text = re.sub(r"</p>|</h1>|</h2>|</div>", "\n", text)
        return re.sub(r"<[^>]+>", "", text)


# ---------------------------------------------------------------------------
# Version snapshot (audit on approval).
# ---------------------------------------------------------------------------
def snapshot_version(
    session: Session, corr: Correspondence, *, render: bool = True
) -> CorrespondenceVersion:
    """Persist a CorrespondenceVersion: rendered html + values (+ pdf/docx bytes).

    Robust: PDF/DOCX rendering failures are logged and swallowed so a snapshot row
    is still written (bytes columns left NULL).
    """
    rendered_html = render_letter_html(session, corr)

    pdf_bytes: Optional[bytes] = None
    docx_bytes: Optional[bytes] = None
    if render:
        try:
            pdf_bytes = render_pdf(session, corr)
        except Exception as exc:  # noqa: BLE001
            logger.warning("snapshot PDF render failed for %s: %s", corr.id, exc)
        try:
            docx_bytes = render_docx(session, corr)
        except Exception as exc:  # noqa: BLE001
            logger.warning("snapshot DOCX render failed for %s: %s", corr.id, exc)

    # Compute next version + INSERT with retry: the UNIQUE(correspondence_id,
    # version) constraint turns a concurrent duplicate into an IntegrityError, on
    # which we recompute the max and try again rather than writing a dup number.
    for attempt in range(5):
        next_version = (
            max(
                (
                    v.version
                    for v in session.exec(
                        select(CorrespondenceVersion).where(
                            CorrespondenceVersion.correspondence_id == corr.id
                        )
                    ).all()
                ),
                default=0,
            )
            + 1
        )
        row = CorrespondenceVersion(
            id=_gen_id("ver"),
            correspondence_id=corr.id,
            version=next_version,
            doc_html=rendered_html,
            values=dict(corr.values or {}),
            pdf_bytes=pdf_bytes,
            docx_bytes=docx_bytes,
            created_at=_now_iso(),
        )
        session.add(row)
        try:
            session.commit()
        except IntegrityError:
            session.rollback()
            if attempt == 4:
                raise
            continue
        session.refresh(row)
        return row
    # Unreachable: the loop returns or raises on the final attempt.
    raise RuntimeError("snapshot_version: exhausted version-number retries")


def snapshot_version_bg(corr_id: str) -> None:
    """Background-task entrypoint: open a FRESH session and snapshot best-effort.

    Runs AFTER the response is sent (the request session is already closed), so it
    manages its own session/engine. Never raises — approval must not be affected.
    """
    try:
        with Session(engine) as session:
            corr = session.get(Correspondence, corr_id)
            if corr is None:
                logger.warning("snapshot_version_bg: correspondence %s gone", corr_id)
                return
            snapshot_version(session, corr)
            logger.info("Version snapshot written for %s", corr_id)
    except Exception as exc:  # noqa: BLE001
        logger.error("snapshot_version_bg failed for %s: %s", corr_id, exc)
